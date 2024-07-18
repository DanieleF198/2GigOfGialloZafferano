import os
import random

import ILASPparser
from io import StringIO
import docx
from docx import oxml
import sys
import numpy as np
import re
import pandas as pd

columns = ["dolcificanti", "farinacei", "erbe_spezie_e_condimenti", "carne", "cereali", "frutta", "funghi_e_tartufi", "latticini", "pasta", "pesce", "uova", "verdure_e_ortaggi", "bollitura", "cottura_a_fiamma", "cottura_a_vapore", "cottura_in_forno", "frittura", "mantecatura", "marinatura", "rosolatura", "stufato", "difficolta", "tempo_di_preparazione", "costo", "antipasti", "piatto_unico", "primo", "secondo", "torta_salata"]

users_of_survey = [10]    # 0, 4, 10, 12, 13, 19, 21, 22, 23, 24, 25, 26, 28, 31, 41, 42, 43, 44

f_preferences_test_local_dir = "./ILASPcode/local/local/Data8Component2Std/sampled-recipes-zero/Train105_gauss/std-0.1/"
fNames = os.path.join('./dataset_100/separated_text_data/names.txt')
fLinks = os.path.join('./dataset_100/separated_text_data/links.txt')

fN = open(fNames)
dataNames = fN.read()
fN.close()
fL = open(fLinks)
dataLinks = fL.read()
fL.close()

linesOfN = dataNames.split('\n')
food_data_names = []
for i, line in enumerate(linesOfN):
    food_data_names.append(line)
linesOfL = dataLinks.split('\n')
food_data_links = []
for i, line in enumerate(linesOfL):
    food_data_links.append(line)

for user in users_of_survey:



    # fPTest = open(os.path.join(f_preferences_test_local_dir + "couple" + str(user) + ".txt"))
    # dataPTest = fPTest.read()
    # fPTest.close()

    # linesOfPTest = dataPTest.split('\n')
    # couples_dataset_test = np.zeros((100, 2), dtype='int32')
    # for i, line in enumerate(linesOfPTest):
    #     if line == "":
    #         continue
    #     couple = [x for x in line.split(';')[0:]]
    #     for j, element in enumerate(couple):
    #         couples_dataset_test[i, j] = int(element)

    doc = docx.Document()

    title = doc.add_heading('Verifica delle preferenze gastronomiche', level=1)

    introduction_paragraph = doc.add_paragraph()
    introduction_paragraph.add_run("\n")
    introduction_paragraph.add_run("Ciao!")
    introduction_paragraph.add_run("\n")
    introduction_paragraph.add_run("In passato le abbiamo chiesto di rispondere ad un sondaggio dove le venivano chieste alcune informazioni riguardo alle vostre preferenze gastronomiche, questo al fine di raccogliere dati da usare in degli studi relativi ad algoritmi di machine learning e sistemi di apprendimento induttivo basati sulla logica (ILASP).")
    introduction_paragraph.add_run("\n")
    introduction_paragraph.add_run("Come forse ricorderai nel sondaggio era richiesta opzionalmente la vostra mail, questo al fine di ricontattarvi e sottoporvi nuovamente ad un sondaggio. Questa volta l'obiettivo è quello di ottenere dati che ci diano delle conferme (o meno) sugli studi condotti, ed avere un feedback su di essi.")
    introduction_paragraph.add_run("\n")
    introduction_paragraph.add_run("Quello che le chiediamo di fare è semplicemente rispondere nuovamente a delle domande riguardanti le sue preferenze gastronomiche. Non richiederà più di 10 minuti e ci permetterà di condurre degli studi più accurati.")
    introduction_paragraph.add_run("\n")
    # introduction_paragraph.add_run("Tutte le ricette presenti sono consultabili sul sito Giallo Zafferano.")
    # introduction_paragraph.add_run("\n")
    introduction_paragraph.add_run("La ringraziamo per il tuo tempo.")
    introduction_paragraph.add_run("\n")


    doc.add_heading("Valutazione delle descrizioni", level=1)
    val_sent_1_intr_paragraph = doc.add_paragraph()
    val_sent_1_intr_paragraph.add_run("\n")
    val_sent_1_intr_paragraph.add_run("Di seguito vengono riportate diverse descrizioni su quelli che potrebbero essere alcuni dei suoi gusti.")
    val_sent_1_intr_paragraph.add_run("\n")
    val_sent_1_intr_paragraph.add_run("Ciascuna descrizione riporta le vostre preferenze in ordine crescente di priorità (dove 1 è la priorità minima mentre 5 è la priorità massima). Se individuate, sono anche riportate le eccezioni, ovvero casi in cui una preferenza viene a mancare a causa della presenza di un certa caratteristica nella relativa ricetta.")
    val_sent_1_intr_paragraph.add_run("\n")

    # buffer = StringIO()
    # sys.stdout = buffer
    # ILASPparser.printTheory('ILASPcode/Data8Component2Std/testOutput/results_zero.csv', user=user, max_v=5, max_p=5, couple=210, data_collector=collect_data, data_counter_collector = counter_collector)
    # sys.stdout = sys.__stdout__
    # translated_theory = ILASPparser.translate_theory(buffer.getvalue())
    # counter_guard = len(translated_theory.split("\n"))
    # val_sent_1 = doc.add_paragraph()
    # for j, line in enumerate(translated_theory.split("\n")):
    #     if j <= 3:
    #         continue
    #     if line == "---------------------------------------------------------------------------------------":
    #         continue
    #     val_sent_1.add_run(line)
    #     if j >= counter_guard - 3:
    #         break
    #     val_sent_1.add_run("\n")
    for subclass_index, subclass in enumerate(["Data", "Data17Component2Std", "Data8Component2Std"]):
        for i in range(5, 6):
            collect_data = pd.DataFrame(np.zeros((50, len(columns)), dtype="float32"), columns=[*columns])
            counter_collector = 0
            buffer = StringIO()
            sys.stdout = buffer
            ILASPparser.printTheory('ILASPcode/' + subclass + '/testOutput/results_zero.csv', user=user, max_v=i, max_p=i, couple=210, data_collector=collect_data, data_counter_collector = counter_collector)
            sys.stdout = sys.__stdout__
            global_theory1 = buffer.getvalue()
            global_theory = ILASPparser.translate_theory(global_theory1)
            doc.add_heading("Descrizione " + str(subclass_index+1), level=1)
            val_sent_1_second_intr = doc.add_paragraph()
            val_sent_1_second_intr.add_run("\n")
            val_sent_1_second_intr.add_run("Date due ricette")

            for j, line in enumerate(global_theory.split("\n")):
                if j <= 3:
                    continue
                if line == "---------------------------------------------------------------------------------------" or line == "-------------------------------------------------- -------------------------------------":
                    continue
                if "priorità" in line:
                    doc.add_paragraph(line, style="ListNumber")
                    val_sent_1 = doc.add_paragraph()
                    val_sent_1.add_run("A. sono d'accordo\n")
                    val_sent_1.add_run("B. sono d'accordo ma gli do maggiore priorità\n")
                    val_sent_1.add_run("C. sono d'accordo ma gli do minore priorità\n")
                    val_sent_1.add_run("D. non sono d'accordo\n")
                if "Seppur sia vero" in line or "Anche se è vero" in line:
                    doc.add_paragraph("eccezione - " + line[2:], style="ListNumber")
                    val_sent_1 = doc.add_paragraph()
                    val_sent_1.add_run("A. sono d'accordo\n")
                    val_sent_1.add_run("B. non sono d'accordo\n")

    # doc.add_heading("Valutazione delle ricette", level=1)
    #
    # question_counter = 1
    #
    # for i, couple in enumerate(couples_dataset_test):
    #     name1 = food_data_names[couple[0]].replace("-", " ").replace("à", "a'")
    #     name2 = food_data_names[couple[1]].replace("-", " ").replace("à", "a'")
    #     link1 = food_data_links[couple[0]]
    #     link2 = food_data_links[couple[1]]
    #     options = doc.add_paragraph()
    #     options.add_run(str(question_counter) + ". Cosa preferisci tra:\n")
    #     question_counter += 1
    #     options.add_run("A. " + name1 + " (" + str(link1) + ")" + "\n")
    #     options.add_run("B. " + name2 + " (" + str(link2) + ")" + "\n")
    #     options.add_run("C. indifferente")
    #

    tempPath = './local-temp-user' + str(user) + '.csv'
    with open(tempPath, 'w+', encoding='UTF8') as f_output:
        f_output.write("USERID;MAXV;MAXP;MAXWC;TRAIN_SIZE;TEST_SIZE;ACCURACYP;PRECISIONP;RECALLP;TRAIN_TIME;THEORY\n")
        filename_couple = './ILASPcode/local/local/Data8Component2Std/sampled-recipes-zero/Train105_gauss/std-0.1/couple' + str(user) + ".txt"
        f_couple = open(filename_couple, "r")
        couple = f_couple.read()
        f_couple.close()
        good_couple = False
        while not good_couple:
            couple_number = random.randint(0, 99)
            for line_index, line in enumerate(couple.split("\n")):
                if line_index == couple_number:
                    first_couple = line
                    break
            for index_element, element in enumerate(first_couple.split(";")):
                if index_element == 0:
                    first_element = element
                elif index_element == 1:
                    second_element = element
                else:
                    break
            filename_theory = 'ILASPcode/local/local/Data8Component2Std/final/users/zero/train/105Couples_gauss_std0.1/User' + str(user) + '/outputTrain/Couple' + first_element + '-' + second_element + '-max_v=1-max_p=5.txt'
            f_local_output = open(filename_theory)
            local_output = f_local_output.read()
            f_local_output.close()
            theory = ""
            for line in local_output.split("\n"):
                if ':~' not in line:
                    continue
                else:
                    theory += line
            if theory != "":
                good_couple = True
        f_output.write(str(user) + ";1;5;3;105;1;0;0;0;0;" + theory + "\n")
    collect_data = pd.DataFrame(np.zeros((50, len(columns)), dtype="float32"), columns=[*columns])
    counter_collector = 0
    buffer = StringIO()
    sys.stdout = buffer
    ILASPparser.printTheory('./local-temp-user' + str(user) + '.csv', user=user, max_v=1, max_p=5, couple=105, data_collector=collect_data, data_counter_collector=counter_collector)
    sys.stdout = sys.__stdout__
    global_theory1 = buffer.getvalue()
    global_theory = ILASPparser.translate_theory(global_theory1)
    doc.add_heading("Descrizione 4", level=1)
    name1 = food_data_names[int(first_element)].replace("-", " ").replace("à", "a'")
    name2 = food_data_names[int(second_element)].replace("-", " ").replace("à", "a'")
    link1 = food_data_links[int(first_element)]
    link2 = food_data_links[int(second_element)]

    val_sent_1_second_intr = doc.add_paragraph()
    val_sent_1_second_intr.add_run("\n")
    val_sent_1_second_intr.add_run("Data la ricetta \"" + str(name1) + "\" (" + str(link1) + ") e la ricetta \"" + str(name2) + "\" (" + str(link2) + ")")

    del val_sent_1

    for j, line in enumerate(global_theory.split("\n")):
        if j <= 3:
            continue
        if line == "---------------------------------------------------------------------------------------" or line == "-------------------------------------------------- -------------------------------------":
            continue
        if "priorità" in line:
            doc.add_paragraph(line, style="ListNumber")
            val_sent_1 = doc.add_paragraph()
            val_sent_1.add_run("A. sono d'accordo\n")
            val_sent_1.add_run("B. sono d'accordo ma gli do maggiore priorità\n")
            val_sent_1.add_run("C. sono d'accordo ma gli do minore priorità\n")
            val_sent_1.add_run("D. non sono d'accordo\n")
        if "Seppur sia vero" in line or "Anche se è vero" in line:
            doc.add_paragraph("eccezione - " + line[2:], style="ListNumber")
            val_sent_1 = doc.add_paragraph()
            val_sent_1.add_run("A. sono d'accordo\n")
            val_sent_1.add_run("B. non sono d'accordo\n")
    doc.add_paragraph("Ritieni che questa descrizione sia pertinente alle due ricette riportate?", style="ListNumber")

    os.remove('./local-temp-user' + str(user) + '.csv')

    doc.add_heading("Domande finali", level=1)
    space_par = doc.add_paragraph()
    space_par.add_run("\n")
    doc.add_paragraph("Ritieni che le descrizioni siano state esposte in maniera chiara?", style="ListNumber")
    doc.add_paragraph("Ritieni che le descrizioni siano state ben scritte?", style="ListNumber")

    doc.save("./survey_of_return/user" + str(user) + ".docx")
    del doc, title, introduction_paragraph, val_sent_1_intr_paragraph, val_sent_1, val_sent_1_second_intr, collect_data, counter_collector
